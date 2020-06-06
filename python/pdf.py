# -*- coding: utf-8 -*-
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
from docx import Document
document = Document()
import warnings
warnings.filterwarnings("ignore")
import os
# 添加你要转换的文件地址
file_name = os.open(r'E:\Download\3G电脑资料大合集\电脑维修内部资料1\[笔记本电脑使用与维护超级技巧1000例].周梁.等编着.影印版.pdf', os.O_RDWR)
def main():
    fn = open(file_name, 'rb')
    parser = PDFParser(fn)
    doc = PDFDocument()
    parser.set_document(doc)
    doc.set_parser(parser)
    resource = PDFResourceManager()
    laparams = LAParams()
    device = PDFPageAggregator(resource, laparams=laparams)
    interpreter = PDFPageInterpreter(resource, device)
    for i in doc.get_pages():
        interpreter.process_page(i)
        layout = device.get_result()
        for out in layout:
            if hasattr(out, "get_text"):
                content = out.get_text().replace(u'\xa0', u' ')
                document.add_paragraph(
                    content, style='ListBullet'
                )
            document.save('a1' + '.docx')
    print('处理完成')
if __name__ == '__main__':
    main()